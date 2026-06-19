import io
import os
import logging
from sqlalchemy.orm import Session

from rag.chunks import (
    clean_text, chunk_text, chunk_by_titles, chunk_by_semantics,
    get_smart_text_end, get_smart_text_start,
    find_sentence_continuation, find_concept_bridge,
)
from rag.embedder import embed_texts
from rag.vectorizer import store_chunks, store_embeddings_batch
from rag.document_processing_utils import deduplicate_chunks

logger = logging.getLogger("uvicorn")

# ------------------------------------------------------------------ #
# EXTRACTION DE TEXTE                                                  #
# ------------------------------------------------------------------ #

def _extract_pdf_pymupdf(content_bytes: bytes) -> list[tuple[int, str]]:
    """Extraction via PyMuPDF avec fallback OCR Tesseract si page vide."""
    import fitz
    pages = []
    doc = fitz.open(stream=content_bytes, filetype="pdf")

    for i, page in enumerate(doc):
        text = page.get_text().strip()

        # Fallback OCR si la page ne contient pas de texte extractible (ex: scan, image)
        if len(text) < 50:
            try:
                import pytesseract
                from PIL import Image
                tesseract_path = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
                # Rendu à 2x pour améliorer la précision OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(img, lang="fra+eng")
            except Exception as e:
                logger.warning(f"OCR page {i+1} échoué : {e}")

        text = clean_text(text)
        if text.strip():
            pages.append((i + 1, text))

    doc.close()
    return pages


def _extract_pdf(content_bytes: bytes) -> list[tuple[int, str]]:
    """Essaie PyMuPDF d'abord, repli sur pypdf."""
    try:
        return _extract_pdf_pymupdf(content_bytes)
    except Exception as e:
        logger.warning(f"PyMuPDF échoué, repli pypdf : {e}")
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content_bytes))
        pages = []
        for i, page in enumerate(reader.pages):
            text = clean_text(page.extract_text() or "")
            if text.strip():
                pages.append((i + 1, text))
        return pages


def _extract_txt(content_bytes: bytes) -> list[tuple[int, str]]:
    # Fichier texte brut : une seule "page"
    text = clean_text(content_bytes.decode("utf-8", errors="replace"))
    return [(1, text)] if text.strip() else []


def _extract_docx(content_bytes: bytes) -> list[tuple[int, str]]:
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(content_bytes))
    parts: list[str] = []

    # Extraction des paragraphes
    for para in doc.paragraphs:
        t = clean_text(para.text)
        if t.strip():
            parts.append(t)

    # Extraction des tableaux (cellules séparées par " | ")
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                clean_text(cell.text) for cell in row.cells if cell.text.strip()
            )
            if row_text.strip():
                parts.append(row_text)

    return [(1, " ".join(parts))] if parts else []


def _extract_xlsx(content_bytes: bytes) -> list[tuple[int, str]]:
    import openpyxl
    # read_only + data_only pour éviter de charger les formules et les styles
    wb = openpyxl.load_workbook(io.BytesIO(content_bytes), read_only=True, data_only=True)
    pages: list[tuple[int, str]] = []

    for sheet_num, sheet in enumerate(wb.worksheets, start=1):
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) for c in row if c is not None and str(c).strip())
            if row_text.strip():
                rows.append(row_text)
        text = clean_text("\n".join(rows))
        if text.strip():
            # Chaque feuille = une page
            pages.append((sheet_num, text))

    wb.close()
    return pages


def _extract_html(content_bytes: bytes) -> list[tuple[int, str]]:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(content_bytes.decode("utf-8", errors="replace"), "html.parser")
    # Supprime les balises non textuelles avant extraction
    for tag in soup(["script", "style", "head", "meta", "link", "noscript"]):
        tag.decompose()
    text = clean_text(soup.get_text(separator=" ", strip=True))
    return [(1, text)] if text.strip() else []


# Dispatch MIME → extracteur ; fallback txt pour tout type inconnu
_MIME_EXTRACTORS = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "application/msword": _extract_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _extract_xlsx,
    "application/vnd.ms-excel": _extract_xlsx,
    "text/html": _extract_html,
}


def extract_pages(content_bytes: bytes, mime_type: str) -> list[tuple[int, str]]:
    extractor = _MIME_EXTRACTORS.get(mime_type)
    if extractor:
        return extractor(content_bytes)
    # Type non reconnu → traité comme texte brut
    return _extract_txt(content_bytes)


# ------------------------------------------------------------------ #
# CHUNKING AVEC CONTEXTE INTER-PAGES                                   #
# ------------------------------------------------------------------ #

OVERLAP_SIZE = 100  # tokens de contexte ajoutés depuis les pages adjacentes


def _build_page_with_context(pages: list[tuple[int, str]], i: int) -> str:
    """Encadre la page i avec la fin de la page précédente et le début de la suivante."""
    _, page_content = pages[i]
    extended = clean_text(page_content)

    # Ajoute la fin de la page précédente pour ne pas perdre le contexte d'entrée
    if i > 0:
        prev_context = get_smart_text_end(clean_text(pages[i - 1][1]), OVERLAP_SIZE)
        if prev_context:
            extended = f"{prev_context}\n\n{extended}"

    # Ajoute le début de la page suivante pour ne pas perdre le contexte de sortie
    if i < len(pages) - 1:
        next_context = get_smart_text_start(clean_text(pages[i + 1][1]), OVERLAP_SIZE)
        if next_context:
            extended = f"{extended}\n\n{next_context}"

    return extended


def _create_transition_chunks(pages: list[tuple[int, str]]) -> list[tuple[int, str]]:
    """Génère des chunks à cheval entre pages consécutives (3 variantes)."""
    chunks = []
    for i in range(len(pages) - 1):
        cur_num, cur_content = pages[i]
        nxt_num, nxt_content = pages[i + 1]
        cur = clean_text(cur_content)
        nxt = clean_text(nxt_content)
        if not cur.strip() or not nxt.strip():
            continue

        # Variante 1 : fin de page courante + début de page suivante
        end = get_smart_text_end(cur, OVERLAP_SIZE)
        start = get_smart_text_start(nxt, OVERLAP_SIZE)
        if end and start:
            transition = f"{end}\n\n{start}"
            chunks.extend(chunk_text(cur_num, transition))

        # Variante 2 : phrase coupée en milieu de saut de page
        continuation = find_sentence_continuation(cur, nxt)
        if continuation:
            chunks.extend(chunk_text(cur_num, continuation))

        # Variante 3 : pont sur les concepts communs aux deux pages
        bridge = find_concept_bridge(cur, nxt, OVERLAP_SIZE)
        if bridge:
            chunks.extend(chunk_text(nxt_num, bridge))

    return chunks


def build_chunks_with_context(pages: list[tuple[int, str]]) -> list[tuple[int, str]]:
    """
    Pipeline complet de chunking :
    - Stratégie 1 : chunks par page avec contexte des pages adjacentes
    - Stratégie 2 : chunks de transition entre pages consécutives
    - Stratégie 3 : chunks par section/titre (title-based)
    """
    all_chunks: list[tuple[int, str]] = []

    # Stratégie 1 : chaque page enrichie du contexte voisin
    for i, (page_num, _) in enumerate(pages):
        extended = _build_page_with_context(pages, i)
        all_chunks.extend(chunk_text(page_num, extended))

    # Stratégie 2 : chunks inter-pages (uniquement si document multi-pages)
    if len(pages) > 1:
        all_chunks.extend(_create_transition_chunks(pages))

    # Stratégie 3 : chunks par section détectée (articles, chapitres, titres Markdown...)
    # Ajoute des chunks structurés uniquement quand des titres sont présents
    for page_num, page_content in pages:
        title_chunks = chunk_by_titles(page_num, clean_text(page_content))
        if len(title_chunks) > 1:  # au moins 2 sections détectées → structure utile
            all_chunks.extend(title_chunks)

    # Stratégie 4 : semantic chunking — coupe aux frontières de changement de sujet
    # Chaque chunk parle d'un seul sujet cohérent (similarité cosinus entre phrases)
    for page_num, page_content in pages:
        semantic_chunks = chunk_by_semantics(page_num, clean_text(page_content))
        if len(semantic_chunks) > 1:  # au moins 2 segments sémantiques détectés
            all_chunks.extend(semantic_chunks)

    return all_chunks


# ------------------------------------------------------------------ #
# PIPELINE COMPLET                                                     #
# ------------------------------------------------------------------ #

def index_document(db: Session, document_id: str, content_bytes: bytes, mime_type: str) -> int:
    pages = extract_pages(content_bytes, mime_type)
    if not pages:
        logger.warning(f"Document {document_id} : aucun texte extrait.")
        return 0

    chunks = build_chunks_with_context(pages)
    if not chunks:
        logger.warning(f"Document {document_id} : aucun chunk généré.")
        return 0

    texts = [c[1] for c in chunks]
    embeddings = embed_texts(texts)
    chunks, embeddings = deduplicate_chunks(chunks, embeddings)

    chunk_ids = store_chunks(db, document_id, chunks)
    store_embeddings_batch(db, list(zip(chunk_ids, embeddings)))

    logger.info(f"Document {document_id} : {len(chunk_ids)} chunks indexés.")
    return len(chunk_ids)
